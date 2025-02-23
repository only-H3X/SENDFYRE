import os
import re
import csv
import random
import asyncio
import logging
import configparser
import ssl
import dns.resolver
import requests
import socket
import math
import io
import base64
import time
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email import encoders
from email.header import Header
from email.utils import formataddr

import aiosmtplib
import socks  # For proxy support
from validate_email_address import validate_email
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF  # For converting TXT/DOCX/HTML to PDF
import qrcode         # For generating QR codes
from PIL import Image  # For converting images to PDF

# Try to import xhtml2pdf for improved HTML-to-PDF conversion.
try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None
    logging.warning("xhtml2pdf module not found. HTML attachments will be converted as plain text.")

# Set logging level to INFO.
logging.getLogger().setLevel(logging.INFO)

# =============================================================================
# Minimal helper definitions to avoid NameError issues
# =============================================================================
def check_email_validity(email):
    regex = r"(^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$)"
    return re.fullmatch(regex, email) is not None

def validate_mx_records(domain):
    try:
        mx_records = dns.resolver.resolve(domain, 'MX')
        return bool(mx_records)
    except Exception:
        return False

# =============================================================================
# GLOBAL VARIABLES AND COUNTERS
# =============================================================================
emails_sent = 0
emails_failed = 0
SMTP_INDEX = 0  # For round-robin SMTP rotation
cc_recipients = []  # Global cc list
bcc_recipients = []  # Global bcc list
DOMAIN_LOGO_CACHE = {}  # Cache for domain logos

# =============================================================================
# RATE LIMITER CLASS
# =============================================================================
class RateLimiter:
    def __init__(self, rate: float):
        self.rate = rate
        self._lock = asyncio.Lock()
        self._last = 0.0

    async def wait(self):
        async with self._lock:
            now = asyncio.get_event_loop().time()
            wait_time = max(0, (1.0 / self.rate) - (now - self._last))
            if wait_time > 0:
                await asyncio.sleep(wait_time)
            self._last = asyncio.get_event_loop().time()

# =============================================================================
# CONFIGURATION & LOGGING SETUP
# =============================================================================
if not os.path.exists("logs"):
    os.makedirs("logs")
LOG_FILE = "logs/email_sender.log"
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
console_handler.setFormatter(console_formatter)
logging.getLogger().addHandler(console_handler)

# -----------------------------------------------------------------------------
# Load external configuration
# -----------------------------------------------------------------------------
CONFIG_FILE = "email_config.ini"
config_parser = configparser.ConfigParser()
if os.path.exists(CONFIG_FILE):
    config_parser.read(CONFIG_FILE, encoding="utf-8")
    logging.info(f"Loaded configuration from {CONFIG_FILE}")
else:
    logging.info("No external config file found. Using default settings.")

def get_config(section, key, default, value_type=str):
    try:
        if value_type == bool:
            return config_parser.getboolean(section, key)
        else:
            return value_type(config_parser.get(section, key))
    except Exception:
        return default

# =============================================================================
# LOAD CONFIGURATION FROM FILE
# =============================================================================
USE_RANDOM_SENDER         = get_config("GENERAL", "USE_RANDOM_SENDER", True, bool)
USE_SENDER_NAME           = get_config("GENERAL", "USE_SENDER_NAME", True, bool)
USE_PROXY                 = get_config("GENERAL", "USE_PROXY", False, bool)
FETCH_DOMAIN_LOGO         = get_config("GENERAL", "FETCH_DOMAIN_LOGO", True, bool)
CHECK_SPF_DKIM_DMARC      = get_config("GENERAL", "CHECK_SPF_DKIM_DMARC", True, bool)
RANDOM_DELAY              = get_config("GENERAL", "RANDOM_DELAY", True, bool)
ENCODE_ATTACHMENTS        = get_config("GENERAL", "ENCODE_ATTACHMENTS", False, bool)
ROTATE_SUBJECTS           = get_config("GENERAL", "ROTATE_SUBJECTS", True, bool)
ATTACHMENTS_TAG_REPLACEMENT = get_config("GENERAL", "ATTACHMENTS_TAG_REPLACEMENT", True, bool)
EMAIL_BODY_TAG_REPLACEMENT = get_config("GENERAL", "EMAIL_BODY_TAG_REPLACEMENT", True, bool)
SEND_DELAY_MIN            = get_config("GENERAL", "SEND_DELAY_MIN", 3, int)
SEND_DELAY_MAX            = get_config("GENERAL", "SEND_DELAY_MAX", 7, int)
MAX_RETRIES               = get_config("GENERAL", "MAX_RETRIES", 3, int)
INITIAL_RETRY_DELAY       = get_config("GENERAL", "INITIAL_RETRY_DELAY", 2, int)
MAX_CONCURRENT_CONNECTIONS = get_config("GENERAL", "MAX_CONCURRENT_CONNECTIONS", 5, int)
SKIP_AUTH                 = get_config("GENERAL", "SKIP_AUTH", False, bool)
USE_TLS                   = get_config("GENERAL", "USE_TLS", True, bool)
ENHANCED_EMAIL_VALIDATION = get_config("GENERAL", "ENHANCED_EMAIL_VALIDATION", True, bool)
EMAILS_PER_SECOND         = get_config("GENERAL", "EMAILS_PER_SECOND", 1.0, float)
BATCH_SENDING             = get_config("GENERAL", "BATCH_SENDING", True, bool)
ROTATE_SMTP               = get_config("GENERAL", "ROTATE_SMTP", False, bool)
RECIPIENTS_FILE           = get_config("GENERAL", "RECIPIENTS_FILE", "recipients.txt")

def load_smtp_nodes():
    nodes_str = get_config("SMTP", "smtp_nodes", "")
    nodes = []
    if nodes_str:
        for entry in nodes_str.split(";"):
            parts = entry.strip().split("|")
            if len(parts) < 4:
                logging.error(f"SMTP node entry '{entry}' does not have enough fields.")
                continue
            node = {
                "server": parts[0].strip(),
                "port": int(parts[1].strip()),
                "email": parts[2].strip(),
                "password": parts[3].strip(),
                "sender_name": parts[4].strip() if len(parts) >= 5 else ""
            }
            nodes.append(node)
    if not nodes:
        logging.error("No SMTP nodes configured in [SMTP] section!")
    return nodes

NODES = load_smtp_nodes()

PROXY_HOST = get_config("PROXY", "HOST", "proxy.example.com")
PROXY_PORT = get_config("PROXY", "PORT", 1080, int)
PROXY_USER = get_config("PROXY", "USER", "")
PROXY_PASS = get_config("PROXY", "PASS", "")

HTML_TEMPLATE_FILE_PATH = get_config("FILES", "HTML_TEMPLATE", "template.html")
attachments = [att.strip() for att in get_config("FILES", "ATTACHMENTS", "sample.txt,sample.pdf,sample.docx").split(",")]

CONVERT_ATTACHMENTS = get_config("CONVERSION", "CONVERT_ATTACHMENTS", True, bool)
CONVERSION_MAPPINGS = get_config("CONVERSION", "CONVERSION_MAPPINGS", "txt:pdf;docx:pdf;jpg:pdf;jpeg:pdf;png:pdf;bmp:pdf;html:pdf", str)
conversion_mappings = {}
if CONVERT_ATTACHMENTS:
    for mapping in CONVERSION_MAPPINGS.split(";"):
        if ":" in mapping:
            src, tgt = mapping.split(":")
            conversion_mappings[src.strip().lower()] = tgt.strip().lower()
SEND_CONVERTED_ATTACHMENT = get_config("CONVERSION", "SEND_CONVERTED_ATTACHMENT", True, bool)
CONVERT_TARGET = get_config("CONVERSION", "CONVERT_TARGET", "converted", str)
CONVERTED_ATTACHMENT_NAME = get_config("CONVERSION", "CONVERTED_ATTACHMENT_NAME", "attachment_doc", str)

ENABLE_QR = get_config("QR", "ENABLE_QR", False, bool)
# Read the list of links from the "Link" key in the QR section.
QR_LINKS = [link.strip() for link in get_config("QR", "Link", "").split(",") if link.strip()]
ROTATION_MODE = get_config("QR", "ROTATION_MODE", "random", str)
QR_CURRENT_INDEX = 0

EMAIL_PRIORITY = get_config("HEADERS", "EMAIL_PRIORITY", None, str)
CUSTOM_HEADER = get_config("HEADERS", "CUSTOM_HEADER", None, str)
ENCODE_HEADERS = get_config("HEADERS", "ENCODE_HEADERS", False, bool)

CUSTOM_FROMMAIL = get_config("SENDER", "CUSTOM_FROMMAIL", "", str)
CUSTOM_SENDER_NAME = get_config("SENDER", "CUSTOM_SENDER_NAME", "", str)

SUBJECTS = []
if config_parser.has_section("SUBJECTS"):
    subjects_str = get_config("SUBJECTS", "subject_lines", "", str)
    if subjects_str:
        SUBJECTS = [s.strip() for s in subjects_str.split(";") if s.strip()]
    else:
        logging.warning("No subject lines defined in [SUBJECTS] section; using defaults.")
        SUBJECTS = [
            "Exclusive Deal for You, {{name}}!",
            "🚀 Special Offer Inside, {{name}}!",
            "Hey {{name}}, Don't Miss Out!"
        ]
else:
    logging.warning("Missing [SUBJECTS] section; using default subjects.")
    SUBJECTS = [
        "Exclusive Deal for You, {{name}}!",
        "🚀 Special Offer Inside, {{name}}!",
        "Hey {{name}}, Don't Miss Out!"
    ]

RECIPIENTS_FILE = get_config("RECIPIENTS", "FILE", "recipients.txt")

DEFAULT_TAGS = {
    "{{date}}": datetime.now().strftime("%Y-%m-%d"),
    "{{time}}": datetime.now().strftime("%H:%M"),
    "{{company}}": "Example Corp",
    "{{website}}": "https://example.com",
    "{{phone}}": "+1 (555) 123-4567",
}

# =============================================================================
# RECIPIENTS LOADING FUNCTION
# =============================================================================
def load_recipients():
    if not os.path.exists(RECIPIENTS_FILE):
        logging.warning(f"Recipients file '{RECIPIENTS_FILE}' not found; using default recipients list.")
        return [
            {"name": "John Doe", "email": "john@example.com", "sender": "user1@node.com"},
            {"name": "Jane Smith", "email": "jane@example.com", "sender": "user2@node.com"},
            {"name": "Bob Jones", "email": "bob@example.com", "sender": "user3@node.com"},
        ]
    ext = os.path.splitext(RECIPIENTS_FILE)[1].lower()
    recipients_list = []
    if ext == ".csv":
        with open(RECIPIENTS_FILE, newline="", encoding="utf-8") as csvfile:
            sample = csvfile.read(1024)
            csvfile.seek(0)
            has_header = csv.Sniffer().has_header(sample)
            csvfile.seek(0)
            if has_header:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    email = row.get("email", "").strip()
                    if not email:
                        continue
                    name = row.get("name", email).strip()
                    sender = row.get("sender", "").strip()
                    recipients_list.append({"name": name, "email": email, "sender": sender})
            else:
                reader = csv.reader(csvfile)
                for row in reader:
                    if not row:
                        continue
                    if len(row) == 1:
                        email = row[0].strip()
                        recipients_list.append({"name": email, "email": email, "sender": ""})
                    elif len(row) == 2:
                        name, email = row[0].strip(), row[1].strip()
                        recipients_list.append({"name": name, "email": email, "sender": ""})
                    else:
                        name, email, sender = row[0].strip(), row[1].strip(), row[2].strip()
                        recipients_list.append({"name": name, "email": email, "sender": sender})
    elif ext == ".txt":
        with open(RECIPIENTS_FILE, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                if "," in line:
                    parts = [part.strip() for part in line.split(",")]
                    if len(parts) == 1:
                        email = parts[0]
                        recipients_list.append({"name": email, "email": email, "sender": ""})
                    elif len(parts) == 2:
                        name, email = parts[0], parts[1]
                        recipients_list.append({"name": name, "email": email, "sender": ""})
                    elif len(parts) >= 3:
                        name, email, sender = parts[0], parts[1], parts[2]
                        recipients_list.append({"name": name, "email": email, "sender": sender})
                else:
                    recipients_list.append({"name": line, "email": line, "sender": ""})
    else:
        logging.warning(f"Unsupported file extension for recipients file: {RECIPIENTS_FILE}")
    return recipients_list

recipients = load_recipients()

def configure_proxy():
    if USE_PROXY:
        try:
            socks.setdefaultproxy(socks.SOCKS5, PROXY_HOST, PROXY_PORT, True, PROXY_USER, PROXY_PASS)
            socket.socket = socks.socksocket
            logging.info(f"Proxy configured: {PROXY_HOST}:{PROXY_PORT}")
        except Exception as e:
            logging.error(f"Error configuring proxy: {e}")

if USE_PROXY:
    configure_proxy()

async def read_file_async(file_path, mode="r", encoding="utf-8"):
    if not os.path.exists(file_path):
        logging.error(f"File not found: {file_path}")
        return None
    return await asyncio.to_thread(lambda: open(file_path, mode, encoding=encoding).read())

# =============================================================================
# CONVERSION FUNCTION (in-memory, with placeholder replacement)
# =============================================================================
def convert_attachment(file_path, target_format, recipient=None):
    base, ext = os.path.splitext(file_path)
    ext = ext.lower().lstrip(".")
    if ext not in conversion_mappings:
        logging.info(f"No conversion mapping for extension '{ext}'; skipping conversion for {file_path}")
        return None
    target = conversion_mappings[ext]
    new_name = f"{CONVERTED_ATTACHMENT_NAME}_{int(time.time())}_{random.randint(1000,9999)}.{target}"
    out_buf = io.BytesIO()
    if target == "pdf":
        if ext in ["txt"]:
            try:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)
                with open(file_path, "r", encoding="utf-8") as f:
                    lines = f.readlines()
                if recipient and ATTACHMENTS_TAG_REPLACEMENT:
                    lines = [replace_tags_in_text(line, recipient) for line in lines]
                for line in lines:
                    pdf.multi_cell(0, 10, txt=line.strip())
                pdf.output(out_buf)
                out_buf.seek(0)
                logging.info(f"Converted {file_path} (txt) to PDF: {new_name}")
                return (new_name, out_buf.getvalue())
            except Exception as e:
                logging.error(f"Error converting {file_path} (txt) to PDF: {e}")
                return None
        elif ext in ["html", "htm"]:
            if pisa:
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        html_content = f.read()
                    if recipient and ATTACHMENTS_TAG_REPLACEMENT:
                        html_content = replace_tags_in_text(html_content, recipient)
                    pisa_status = pisa.CreatePDF(html_content, dest=out_buf)
                    if pisa_status.err:
                        logging.error(f"Error converting HTML {file_path} to PDF using xhtml2pdf.")
                        return None
                    out_buf.seek(0)
                    logging.info(f"Converted {file_path} (html) to PDF using xhtml2pdf: {new_name}")
                    return (new_name, out_buf.getvalue())
                except Exception as e:
                    logging.error(f"Error converting HTML {file_path} to PDF: {e}")
                    return None
            else:
                try:
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.set_font("Arial", size=12)
                    with open(file_path, "r", encoding="utf-8") as f:
                        lines = f.readlines()
                    if recipient and ATTACHMENTS_TAG_REPLACEMENT:
                        lines = [replace_tags_in_text(line, recipient) for line in lines]
                    for line in lines:
                        pdf.multi_cell(0, 10, txt=line.strip())
                    pdf.output(out_buf)
                    out_buf.seek(0)
                    logging.info(f"Converted {file_path} (html) to PDF as plain text: {new_name}")
                    return (new_name, out_buf.getvalue())
                except Exception as e:
                    logging.error(f"Error converting HTML {file_path} as plain text to PDF: {e}")
                    return None
        elif ext == "docx":
            try:
                doc = Document(file_path)
                if recipient and ATTACHMENTS_TAG_REPLACEMENT:
                    for para in doc.paragraphs:
                        para.text = replace_tags_in_text(para.text, recipient)
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        pdf.multi_cell(0, 10, txt=text)
                pdf.output(out_buf)
                out_buf.seek(0)
                logging.info(f"Converted {file_path} (docx) to PDF: {new_name}")
                return (new_name, out_buf.getvalue())
            except Exception as e:
                logging.error(f"Error converting {file_path} (docx) to PDF: {e}")
                return None
        elif ext in ["jpg", "jpeg", "png", "bmp"]:
            try:
                image = Image.open(file_path)
                if image.mode != "RGB":
                    image = image.convert("RGB")
                image.save(out_buf, "PDF")
                out_buf.seek(0)
                logging.info(f"Converted {file_path} (image) to PDF: {new_name}")
                return (new_name, out_buf.getvalue())
            except Exception as e:
                logging.error(f"Error converting image {file_path} to PDF: {e}")
                return None
        else:
            logging.info(f"No conversion logic for {file_path} with extension {ext}; skipping conversion")
            return None
    else:
        logging.info(f"Target conversion format '{target}' not supported; skipping conversion for {file_path}")
        return None

# =============================================================================
# READ FILE INTO MEMORY (for non-converted attachments)
# =============================================================================
def read_file_into_memory(file_path):
    try:
        with open(file_path, "rb") as f:
            return f.read()
    except Exception as e:
        logging.error(f"Error reading file {file_path} into memory: {e}")
        return None

# =============================================================================
# QR CODE GENERATION FUNCTION (in-memory)
# =============================================================================
def generate_qr_code_image(link):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf.getvalue()

def get_qr_link():
    global QR_CURRENT_INDEX
    if ROTATION_MODE.lower() == "random":
        return random.choice(QR_LINKS)
    elif ROTATION_MODE.lower() == "sequential":
        link = QR_LINKS[QR_CURRENT_INDEX]
        QR_CURRENT_INDEX = (QR_CURRENT_INDEX + 1) % len(QR_LINKS)
        return link
    else:
        return random.choice(QR_LINKS)

# =============================================================================
# PERSONALIZATION FUNCTIONS
# =============================================================================
def replace_tags_in_text(content, recipient):
    if not EMAIL_BODY_TAG_REPLACEMENT:
        return content
    tags = {**DEFAULT_TAGS, "{{name}}": recipient["name"], "{{email}}": recipient["email"]}
    for tag, value in tags.items():
        content = content.replace(tag, value)
    # Replace QR code placeholder with a data URI
    if "##qrcode##" in content:
        # Generate a QR code image for the given link (if available via recipient context, you may extend this)
        # Here we use a default link if none is provided.
        default_link = "https://example.com"
        qr_data = generate_qr_code_image(default_link)
        data_uri = f"data:image/png;base64,{base64.b64encode(qr_data).decode('utf-8')}"
        content = content.replace("##qrcode##", f'<img src="{data_uri}" alt="QR Code">')
    # Replace domain logo placeholder.
    if "##domainlogo##" in content:
        domain = recipient["email"].split("@")[-1]
        logo = get_domain_logo(domain)
        if logo:
            content = content.replace("##domainlogo##", f'<img src="data:image/png;base64,{logo}" alt="Domain Logo">')
        else:
            content = content.replace("##domainlogo##", "")
    return content

def replace_tags_in_docx(file_path, recipient):
    if not ATTACHMENTS_TAG_REPLACEMENT:
        return None
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            para.text = replace_tags_in_text(para.text, recipient)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_{recipient['email']}.docx"
        return (filename, buf.getvalue())
    except Exception as e:
        logging.error(f"Error processing DOCX {file_path}: {e}")
        return None

def replace_tags_in_txt(file_path, recipient):
    if not ATTACHMENTS_TAG_REPLACEMENT:
        return None
    if not os.path.exists(file_path):
        logging.error(f"File not found: {file_path}")
        return None
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        new_content = replace_tags_in_text(content, recipient)
        filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_{recipient['email']}.txt"
        return (filename, new_content.encode("utf-8"))
    except Exception as e:
        logging.error(f"Error processing TXT {file_path}: {e}")
        return None

def replace_tags_in_pdf(file_path, recipient):
    if not ATTACHMENTS_TAG_REPLACEMENT:
        return None
    if not os.path.exists(file_path):
        logging.error(f"File not found: {file_path}")
        return None
    try:
        reader = PdfReader(file_path)
        writer = PdfWriter()
        for page in reader.pages:
            text = page.extract_text()
            if text:
                replaced_text = replace_tags_in_text(text, recipient)
            writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_{recipient['email']}.pdf"
        return (filename, buf.getvalue())
    except Exception as e:
        logging.error(f"Error processing PDF {file_path}: {e}")
        return None

def personalize_attachments(recipient):
    attachments_list = []
    for file in attachments:
        file = file.strip()
        if file.lower().startswith("http://") or file.lower().startswith("https://"):
            if ENABLE_QR:
                try:
                    qr_data = generate_qr_code_image(file)
                    filename = f"qr_attachment_{recipient['email'].replace('@','_')}_{abs(hash(file))}.png"
                    attachments_list.append((filename, qr_data))
                except Exception as e:
                    logging.error(f"Error generating QR code for URL {file}: {e}")
            else:
                logging.info(f"Skipping URL attachment {file} as QR generation is disabled.")
            continue
        if os.path.exists(file):
            replaced_file = None
            ext = os.path.splitext(file)[1].lower().lstrip(".")
            if ATTACHMENTS_TAG_REPLACEMENT:
                if ext == "docx":
                    replaced_file = replace_tags_in_docx(file, recipient)
                elif ext == "txt":
                    replaced_file = replace_tags_in_txt(file, recipient)
                elif ext == "pdf":
                    replaced_file = replace_tags_in_pdf(file, recipient)
            if replaced_file:
                processed_file = replaced_file
            else:
                if CONVERT_ATTACHMENTS and ext in conversion_mappings:
                    conversion_result = convert_attachment(file, conversion_mappings[ext], recipient)
                    if conversion_result:
                        processed_file = conversion_result
                    else:
                        processed_file = (os.path.basename(file), read_file_into_memory(file))
                else:
                    processed_file = (os.path.basename(file), read_file_into_memory(file))
        else:
            logging.error(f"Attachment file not found: {file}. Skipping.")
            continue
        if SEND_CONVERTED_ATTACHMENT:
            if CONVERT_TARGET.lower() == "both":
                original_data = read_file_into_memory(file)
                attachments_list.append((os.path.basename(file), original_data))
                if processed_file[0] != os.path.basename(file):
                    attachments_list.append(processed_file)
            elif CONVERT_TARGET.lower() == "converted":
                attachments_list.append(processed_file)
            else:
                attachments_list.append((os.path.basename(file), read_file_into_memory(file)))
        else:
            attachments_list.append((os.path.basename(file), read_file_into_memory(file)))
    return attachments_list

def build_from_header(sender_email, sender_name):
    if CUSTOM_FROMMAIL:
        sender_email = CUSTOM_FROMMAIL
    if CUSTOM_SENDER_NAME:
        sender_name = CUSTOM_SENDER_NAME
    if USE_SENDER_NAME and sender_name:
        if ENCODE_HEADERS:
            encoded_name = str(Header(sender_name, "utf-8"))
            return formataddr((encoded_name, sender_email))
        else:
            return f"{sender_name} <{sender_email}>"
    else:
        return sender_email

def get_smtp_node(recipient):
    global SMTP_INDEX
    if ROTATE_SMTP:
        node = NODES[SMTP_INDEX % len(NODES)]
        SMTP_INDEX += 1
        return node
    elif USE_RANDOM_SENDER:
        return random.choice(NODES)
    else:
        sender_email = recipient.get("sender")
        return next((node for node in NODES if node["email"] == sender_email), NODES[0])

def get_subject_line(recipient):
    if ROTATE_SUBJECTS and SUBJECTS:
        raw_subject = random.choice(SUBJECTS)
        subject_text = replace_tags_in_text(raw_subject, recipient)
    else:
        subject_text = f"Hello {recipient['name']}, Your Personalized Offer!"
    if ENCODE_HEADERS:
        return str(Header(subject_text, "utf-8"))
    else:
        return subject_text

def get_domain_logo(domain):
    if not FETCH_DOMAIN_LOGO:
        return None
    if domain in DOMAIN_LOGO_CACHE:
        return DOMAIN_LOGO_CACHE[domain]
    logo_url = f"https://logo.clearbit.com/{domain}"
    try:
        response = requests.get(logo_url, timeout=5)
        if response.status_code == 200:
            buf = io.BytesIO(response.content)
            buf.seek(0)
            encoded_logo = base64.b64encode(buf.getvalue()).decode()
            DOMAIN_LOGO_CACHE[domain] = encoded_logo
            return encoded_logo
    except requests.RequestException as e:
        logging.warning(f"Could not fetch logo for {domain}: {e}")
    return None

# =============================================================================
# CUSTOM PLACEHOLDERS (for advanced personalization)
# =============================================================================
def process_string(value):
    """Removes special characters and converts to lowercase."""
    return re.sub(r'[^a-zA-Z0-9]', '', value).lower()

def encode_base64_custom(value, no_padding=False):
    encoded = base64.b64encode(value.encode()).decode()
    return encoded.rstrip('=') if no_padding else encoded

def encode_hex_custom(value):
    return value.encode().hex()

def obfuscate_hex_custom(value):
    return ''.join(f"{ord(c):x}" for c in value)

def replace_placeholders(content, recipient, sender_name, sender_domain, link):
    """Replaces custom placeholders in the email body and subject."""
    victim_name = process_string(recipient["name"])
    victim_email = recipient["email"]
    victim_domain = process_string(recipient["email"].split("@")[-1])
    victim_full_domain = recipient["email"].split("@")[-1]

    placeholders = {
        "##victimname##": victim_name,
        "##victimemail##": victim_email,
        "##victimdomain##": victim_domain,
        "##victimfulldomain##": victim_full_domain,
        "##victimb64email##": encode_base64_custom(victim_email),
        "##victimb64emailnp##": encode_base64_custom(victim_email, no_padding=True),
        "##victimhexemail##": encode_hex_custom(victim_email),
        "##victimobfhexemail##": obfuscate_hex_custom(victim_email),
        "##victimb64domain##": encode_base64_custom(victim_domain),
        "##victimb64name##": encode_base64_custom(victim_name),
        "##victimdomainlogo##": f'<img src="data:image/png;base64,{get_domain_logo(victim_domain)}" alt="Domain Logo">' if get_domain_logo(victim_domain) else "",
        "##victimdomainlogosrc##": get_domain_logo(victim_domain) or "",
        "##victimrealdomain##": victim_full_domain,
        "##myname##": sender_name,
        "##mydomain##": sender_domain,
        "##link##": link,
        "##linkb64##": encode_base64_custom(link),
        "##qrcode##": f'<img src="data:image/png;base64,{base64.b64encode(generate_qr_code_image(link)).decode("utf-8")}" alt="QR Code">',
        "##qrcodedata##": encode_base64_custom(generate_qr_code_image(link).decode("latin1")),
        "##domainlogo##": f'<img src="data:image/png;base64,{get_domain_logo(victim_domain)}" alt="Domain Logo">' if get_domain_logo(victim_domain) else "",
        "##date(1)##": datetime.now().strftime("%Y-%m-%d"),
    }

    for placeholder, value in placeholders.items():
        content = content.replace(placeholder, value)

    return content

# =============================================================================
# EMAIL SENDER CLASS
# =============================================================================
class EmailSender:
    def __init__(self, semaphore, rate_limiter):
        self.semaphore = semaphore
        self.rate_limiter = rate_limiter

    async def send_email(self, recipient):
        global emails_sent, emails_failed
        if RANDOM_DELAY:
            delay = random.randint(SEND_DELAY_MIN, SEND_DELAY_MAX)
            logging.info(f"Delaying email to {recipient['email']} for {delay} seconds")
            await asyncio.sleep(delay)
        if not check_email_validity(recipient["email"]):
            logging.error(f"Skipping invalid email: {recipient['email']}")
            emails_failed += 1
            return
        domain = recipient["email"].split("@")[-1]
        if not validate_mx_records(domain):
            logging.error(f"Domain {domain} cannot receive emails. Skipping {recipient['email']}.")
            emails_failed += 1
            return
        node = get_smtp_node(recipient)
        sender_email = node["email"]
        sender_name = node["sender_name"]
        logo_data = get_domain_logo(domain)
        html_template = await read_file_async(HTML_TEMPLATE_FILE_PATH)
        if html_template is None:
            logging.error("HTML template not found. Aborting email send.")
            emails_failed += 1
            return
        # Process standard tag replacements.
        html_content = replace_tags_in_text(html_template, recipient)
        # Choose a QR link from the config (if enabled) using rotation.
        if ENABLE_QR and QR_LINKS:
            chosen_qr_link = get_qr_link()
        else:
            chosen_qr_link = ""
        # Apply custom placeholder replacements (including ##qrcode## and ##domainlogo##).
        html_content = replace_placeholders(html_content, recipient, sender_name, domain, chosen_qr_link)
        msg = MIMEMultipart()
        msg["From"] = build_from_header(sender_email, sender_name)
        msg["To"] = recipient["email"]
        msg["Subject"] = get_subject_line(recipient)
        if EMAIL_PRIORITY:
            msg["X-Priority"] = EMAIL_PRIORITY
            if EMAIL_PRIORITY == "1":
                msg["Priority"] = "urgent"
            elif EMAIL_PRIORITY == "5":
                msg["Priority"] = "non-urgent"
        if CUSTOM_HEADER:
            if ENCODE_HEADERS:
                msg["X-Custom-Header"] = str(Header(CUSTOM_HEADER, "utf-8"))
            else:
                msg["X-Custom-Header"] = CUSTOM_HEADER
        msg.attach(MIMEText(html_content, "html"))
        if FETCH_DOMAIN_LOGO and logo_data:
            try:
                # Attach logo from memory (decoded from base64)
                logo_attachment = MIMEImage(base64.b64decode(logo_data), _subtype="png")
                logo_attachment.add_header("Content-ID", "<logo_image>")
                logo_attachment.add_header("Content-Disposition", "inline", filename="logo.png")
                msg.attach(logo_attachment)
            except Exception as e:
                logging.error(f"Error attaching logo: {e}")
        if ENABLE_QR and QR_LINKS:
            try:
                # Use the same chosen QR link for embedding.
                qr_data = generate_qr_code_image(chosen_qr_link)
                qr_attachment = MIMEImage(qr_data, _subtype="png")
                qr_attachment.add_header("Content-ID", "<qr_code>")
                qr_attachment.add_header("Content-Disposition", "inline", filename="qr_code.png")
                msg.attach(qr_attachment)
                logging.info(f"QR Code generated and attached for link: {chosen_qr_link}")
            except Exception as e:
                logging.error(f"Error generating or attaching inline QR code: {e}")
        attachments_data = personalize_attachments(recipient)
        for filename, data in attachments_data:
            if data:
                try:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(data)
                    encoders.encode_base64(part)
                    part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
                    msg.attach(part)
                except Exception as e:
                    logging.error(f"Error attaching in-memory file {filename}: {e}")
        all_recipients = [recipient["email"]] + cc_recipients + bcc_recipients
        await self.rate_limiter.wait()
        attempt = 0
        while attempt < MAX_RETRIES:
            try:
                async with self.semaphore:
                    context = ssl.create_default_context() if USE_TLS else None
                    async with aiosmtplib.SMTP(
                        hostname=node["server"],
                        port=node["port"],
                        start_tls=USE_TLS,
                        tls_context=context
                    ) as server:
                        if not SKIP_AUTH:
                            await server.login(sender_email, node["password"])
                        await server.send_message(msg, sender=sender_email, recipients=all_recipients)
                logging.info(f"✅ Email sent to {recipient['email']} from {sender_email}")
                emails_sent += 1
                break
            except Exception as e:
                attempt += 1
                delay = INITIAL_RETRY_DELAY * math.pow(2, attempt - 1)
                logging.error(f"❌ Attempt {attempt} failed to send email to {recipient['email']}: {e}. Retrying in {delay} seconds.")
                await asyncio.sleep(delay)
        else:
            logging.error(f"❌ All retry attempts failed for {recipient['email']}.")
            emails_failed += 1

async def send_bulk_emails():
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CONNECTIONS)
    rate_limiter = RateLimiter(EMAILS_PER_SECOND)
    sender = EmailSender(semaphore, rate_limiter)
    tasks = [sender.send_email(recipient) for recipient in recipients]
    await asyncio.gather(*tasks)
    total = len(recipients)
    logging.info(f"Email Summary: Attempted: {total}, Sent: {emails_sent}, Failed: {emails_failed}")

async def send_emails_sequentially():
    semaphore = asyncio.Semaphore(1)
    rate_limiter = RateLimiter(EMAILS_PER_SECOND)
    sender = EmailSender(semaphore, rate_limiter)
    for recipient in recipients:
        await sender.send_email(recipient)
    total = len(recipients)
    logging.info(f"Email Summary: Attempted: {total}, Sent: {emails_sent}, Failed: {emails_failed}")

if __name__ == "__main__":
    total_emails = len(recipients)
    num_proxies = 1 if USE_PROXY else 0
    num_subjects = len(SUBJECTS)
    num_frommails = 1 if CUSTOM_FROMMAIL else len(NODES)
    num_smtps = len(NODES)
    logging.info(f"Loaded {total_emails} emails to send")
    logging.info(f"Proxy in use: {'Yes' if USE_PROXY else 'No'} ({num_proxies} configured)")
    logging.info(f"Loaded {num_subjects} subject(s)")
    logging.info(f"Loaded {num_frommails} frommail(s)")
    logging.info(f"Loaded {num_smtps} SMTP node(s)")
    logging.info(f"Send rate: {EMAILS_PER_SECOND} email(s) per second")
    logging.info(f"Batch sending is {'enabled' if BATCH_SENDING else 'disabled'}")
    logging.info(f"SMTP rotation is {'enabled' if ROTATE_SMTP else 'disabled'}")
    try:
        if BATCH_SENDING:
            asyncio.run(send_bulk_emails())
        else:
            asyncio.run(send_emails_sequentially())
    except Exception as main_exc:
        logging.error(f"Critical error in email processing: {main_exc}")
